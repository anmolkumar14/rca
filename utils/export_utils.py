import os
from pathlib import Path

from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

EXPORT_DIR = Path(__file__).resolve().parent.parent / "exports"
EXPORT_DIR.mkdir(exist_ok=True)


def export_document(document_text, summary_text, format_name):
    timestamp = __import__('datetime').datetime.now().strftime("%Y%m%d-%H%M%S")
    safe_name = "rca-" + timestamp

    if format_name.lower() == "pdf":
        file_path = EXPORT_DIR / f"{safe_name}.pdf"
        build_pdf(file_path, document_text, summary_text)
    else:
        file_path = EXPORT_DIR / f"{safe_name}.docx"
        build_word(file_path, document_text, summary_text)

    return str(file_path)


def build_pdf(file_path, document_text, summary_text):
    doc = SimpleDocTemplate(str(file_path), pagesize=letter)
    styles = getSampleStyleSheet()
    story = [Paragraph("AI RCA Generator", styles['Title']), Spacer(1, 12),
             Paragraph("Executive Summary", styles['Heading2']),
             Paragraph(summary_text, styles['BodyText']), Spacer(1, 12),
             Paragraph("Full RCA", styles['Heading2'])]

    for line in document_text.splitlines():
        story.append(Paragraph(line, styles['BodyText']))
        story.append(Spacer(1, 6))

    doc.build(story)


def build_word(file_path, document_text, summary_text):
    doc = Document()
    doc.add_heading('AI RCA Generator', level=1)
    doc.add_heading('Executive Summary', level=2)
    doc.add_paragraph(summary_text)
    doc.add_heading('Full RCA', level=2)
    doc.add_paragraph(document_text)
    doc.save(file_path)
